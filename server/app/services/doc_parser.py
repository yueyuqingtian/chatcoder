"""文件解析服务：图片 base64、CSV/Excel/TSV 表格、Markdown/文本/JSON 等。

用于消息输入框附件上传 — 将文件转为 LLM 可理解的文本或 data URL。
"""
import base64
import csv
import io
import json
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# ── 支持的文件类型 ──
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg"}
SPREADSHEET_EXTS = {".csv", ".tsv", ".xlsx", ".xls"}
DOC_EXTS = {".docx", ".pdf"}
TEXT_EXTS = {".txt", ".md", ".json", ".yaml", ".yml", ".xml", ".html", ".htm", ".log", ".py", ".js", ".ts", ".tsx", ".jsx", ".java", ".go", ".rs", ".c", ".cpp", ".h", ".sh", ".sql"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_SPREADSHEET_ROWS = 200  # 表格最大预览行数


def is_image(filename: str) -> bool:
    return Path(filename).suffix.lower() in IMAGE_EXTS


def is_spreadsheet(filename: str) -> bool:
    return Path(filename).suffix.lower() in SPREADSHEET_EXTS


def is_document(filename: str) -> bool:
    return Path(filename).suffix.lower() in DOC_EXTS


def is_text(filename: str) -> bool:
    return Path(filename).suffix.lower() in TEXT_EXTS


def _image_to_data_url(data: bytes, suffix: str) -> str:
    """将图片字节转为 data URL。"""
    ext = suffix.lower().lstrip(".")
    if ext == "jpg":
        ext = "jpeg"
    b64 = base64.b64encode(data).decode("ascii")
    return f"data:image/{ext};base64,{b64}"


def _parse_csv_like(data: bytes, delimiter: str = ",") -> str:
    """解析 CSV/TSV 为 Markdown 表格文本。"""
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = list(reader)[:MAX_SPREADSHEET_ROWS]
    if not rows:
        return "(空文件)"
    # 构造 Markdown 表格
    header = rows[0]
    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    for row in rows[1:]:
        # 补齐列数
        while len(row) < len(header):
            row.append("")
        lines.append("| " + " | ".join(row) + " |")
    if len(list(reader)) > MAX_SPREADSHEET_ROWS:
        lines.append(f"\n(仅显示前 {MAX_SPREADSHEET_ROWS} 行)")
    return "\n".join(lines)


def _parse_excel(data: bytes) -> str:
    """解析 .xlsx/.xls 为 Markdown 表格文本。"""
    try:
        import openpyxl  # type: ignore
    except ImportError:
        return "(需要安装 openpyxl 才能解析 Excel 文件)"

    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as e:
        logger.warning("[doc_parser] Excel 解析失败: %s", e)
        return f"(Excel 解析失败: {e})"

    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"### Sheet: {sheet_name}\n")
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            if row_count >= MAX_SPREADSHEET_ROWS:
                parts.append(f"(仅显示前 {MAX_SPREADSHEET_ROWS} 行)")
                break
            cells = [str(c) if c is not None else "" for c in row]
            if row_count == 0:
                parts.append("| " + " | ".join(cells) + " |")
                parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
            else:
                parts.append("| " + " | ".join(cells) + " |")
            row_count += 1
        parts.append("")
    wb.close()
    return "\n".join(parts) if parts else "(空 Excel)"


def _parse_docx(data: bytes) -> str:
    """解析 .docx 为纯文本（python-docx）。"""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return "(需要安装 python-docx 才能解析 Word 文档)"

    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # 表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts) if parts else "(空 Word 文档)"
    except Exception as e:
        logger.warning("[doc_parser] docx 解析失败: %s", e)
        return f"(Word 解析失败: {e})"


def _parse_pdf(data: bytes) -> str:
    """解析 .pdf 为纯文本（pypdf）。"""
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        return "(需要安装 pypdf 才能解析 PDF 文档)"

    try:
        reader = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"--- 第 {i + 1} 页 ---\n{text.strip()}")
        return "\n\n".join(parts) if parts else "(PDF 无可提取文本，可能是扫描件/纯图片)"
    except Exception as e:
        logger.warning("[doc_parser] pdf 解析失败: %s", e)
        return f"(PDF 解析失败: {e})"


def parse_file(filename: str, data: bytes) -> dict:
    """解析上传文件，返回 {type, content, data_url, mime}。

    type: "image" | "text" | "spreadsheet" | "document" | "unsupported"
    - image: content 为 data_url
    - text/spreadsheet/document: content 为解析后的文本
    """
    suffix = Path(filename).suffix
    size = len(data)

    if size > MAX_FILE_SIZE:
        return {
            "type": "unsupported",
            "content": f"(文件过大: {size // 1024}KB, 超过 {MAX_FILE_SIZE // 1024 // 1024}MB 限制)",
            "data_url": None,
            "mime": None,
        }

    if is_image(filename):
        data_url = _image_to_data_url(data, suffix)
        return {
            "type": "image",
            "content": data_url,
            "data_url": data_url,
            "mime": f"image/{suffix.lstrip('.').lower().replace('jpg', 'jpeg')}",
        }

    if is_spreadsheet(filename):
        if suffix.lower() in (".csv",):
            text = _parse_csv_like(data, ",")
        elif suffix.lower() == ".tsv":
            text = _parse_csv_like(data, "\t")
        else:
            text = _parse_excel(data)
        return {
            "type": "spreadsheet",
            "content": f"📎 **{filename}**\n\n{text}",
            "data_url": None,
            "mime": "text/markdown",
        }

    if is_document(filename):
        if suffix.lower() == ".docx":
            text = _parse_docx(data)
        else:
            text = _parse_pdf(data)
        return {
            "type": "document",
            "content": f"📎 **{filename}**\n\n{text}",
            "data_url": None,
            "mime": "text/plain",
        }

    if is_text(filename):
        text = data.decode("utf-8", errors="replace")
        return {
            "type": "text",
            "content": f"📎 **{filename}**\n```\n{text}\n```",
            "data_url": None,
            "mime": "text/plain",
        }

    return {
        "type": "unsupported",
        "content": f"(不支持的文件类型: {suffix})",
        "data_url": None,
        "mime": None,
    }


def extract_plain_text(filename: str, data: bytes) -> str | None:
    """提取文件的纯文本内容（供 AI 工具阅读或上下文注入）。

    图片返回 None（AI 应通过 view_image 读取）；
    文本/表格/文档返回提取后的文本；不支持的返回 None。
    """
    parsed = parse_file(filename, data)
    if parsed["type"] == "unsupported" or parsed["type"] == "image":
        return None
    content = parsed.get("content") or ""
    # 去掉 parse_file 拼装的 📎 **filename** 前缀（markdown 代码围栏一并去除）
    if content.startswith("📎 **"):
        lines = content.split("\n", 1)
        content = lines[1] if len(lines) > 1 else ""
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
    return content.strip()
